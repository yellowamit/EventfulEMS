from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("Ashika Final Report - Filled.docx")
OUT_TEXT = Path("Ashika Final Report Expanded Text.txt")


EXPANSIONS = {
    "CHAPTER 1 - INTRODUCTION": [
        "The purpose of this report is to present my training work in a clean and practical manner. I have kept the language simple so that the reader can understand the project flow without needing to inspect the source code. The report also explains why each module was included and how it contributes to the complete event management workflow.",
        "During the training, I tried to think like both a developer and a data analyst. As a developer, I focused on building features that work correctly. As a data analyst, I focused on what useful information the application can collect from events, bookings, ticket quantities, revenue values, and user activity.",
    ],
    "1.1 About Sabudh Foundation": [
        "Sabudh Foundation creates an environment where students are encouraged to learn by doing. The learning style is not limited to theory or classroom-style notes. Instead, interns are expected to explore problems, understand the domain, discuss approaches, and then implement working solutions.",
        "This type of training was helpful for me because modern computer science work often requires more than one skill at the same time. A project may need frontend development, backend APIs, database design, analytics thinking, debugging, deployment awareness, and clear documentation. Sabudh's project-based environment supported that kind of combined learning.",
    ],
    "1.2 Introduction to Internship": [
        "The internship started with the goal of improving my practical understanding of data science and software systems. While working on assignments and project ideas, I understood that data science is not only about models and algorithms. It also depends on how data is collected, stored, cleaned, and made available for analysis.",
        "EventfulEMS became a suitable project for this training because events naturally generate useful data. Every event has a date, time, location, organizer, price, capacity, audience response, and bookings. If these values are stored properly, the same system can later support reports, dashboards, and decision-making.",
    ],
    "1.3 Internship Learning Experience": [
        "I also learned the importance of writing code in a way that can be maintained later. Small decisions such as naming API routes clearly, keeping user state in one context, validating backend requests, and storing configuration in environment variables made the system easier to understand.",
        "Another important learning was that errors should be handled politely. Instead of letting the user face a blank screen, the application should show clear messages such as login required, ticket limit reached, event not found, or not enough tickets available. These small details improve the quality of the final project.",
    ],
    "1.4 Vision Language Navigation Work": [
        "This work was kept as a short research-oriented task in the training. It helped me understand how language, computer vision, and path planning can be combined in robotics. The project was not the main implementation work of this report, but it improved my understanding of multi-modal AI systems.",
        "The most important idea I learned from VLN was that an AI agent must connect words with actual surroundings. A sentence like 'turn left near the sofa and move towards the hallway' is simple for humans, but difficult for a robot because it must recognize objects, understand direction, avoid repeated paths, and judge whether it has reached the correct destination.",
    ],
    "CHAPTER 2 - TECHNOLOGY USED": [
        "The technology stack was selected to support a complete web application. React handled the user interface, Express handled server-side routes, MongoDB handled data storage, and supporting packages handled authentication, image uploads, date formatting, icons, QR generation, and deployment.",
        "I tried to use each tool for a clear reason. Instead of adding many libraries without purpose, the project uses tools that directly support the application's features. This made the project easier to explain, test, and improve.",
    ],
    "2.1 React and Vite": [
        "React's component-based approach was useful because the same style of layout appears in multiple places. For example, event cards, buttons, confirmation boxes, and page containers can be understood as reusable interface parts. This made the frontend easier to grow as new pages were added.",
        "Vite was helpful during development because it starts quickly and refreshes changes fast. This was especially useful while working on forms, event listing, and booking pages, where small changes had to be checked repeatedly in the browser.",
    ],
    "2.2 Tailwind CSS and Frontend Styling": [
        "Tailwind CSS allowed styling to be written close to the JSX markup. This was useful during fast development because I could adjust spacing, colors, borders, and responsive layouts directly inside components. It also helped keep the interface consistent across different pages.",
        "The styling goal was to keep the application simple and usable. The project does not try to look like a marketing website. It focuses on the actual workflow: see events, understand details, create events, book tickets, and view records.",
    ],
    "2.3 Node.js and Express": [
        "Express was used as the central layer between frontend requests and database operations. It receives requests from React, validates required data, performs database queries through Mongoose, and returns responses in JSON format.",
        "The backend also contains important business logic. For example, ticket booking is not only a database insert. It checks user and event IDs, verifies event existence, checks maximum booking limit, checks event capacity, generates a ticket code, saves the ticket, and updates the event count.",
    ],
    "2.4 MongoDB and Mongoose": [
        "MongoDB was convenient for this project because events and tickets can contain nested information. A ticket record stores ticket details such as event name, date, time, price, total price, QR code data, and user details in one document-like structure.",
        "Mongoose improved reliability by defining schemas. Even though MongoDB is flexible, using schemas made it easier to maintain expected fields and validations such as required ticket details, minimum ticket price, ticket status values, and quantity limits.",
    ],
    "2.5 Authentication and Security Tools": [
        "Authentication was important because the application needs to know who created an event and who booked a ticket. Without authentication, ticket records and event ownership would not be reliable.",
        "The use of HTTP-only cookies improves security because the token is not directly available to frontend JavaScript. Password hashing with bcryptjs also protects users because even the database does not store their original password.",
    ],
    "2.6 API, Upload, and QR Tools": [
        "Image upload support made event pages more realistic because events are usually promoted visually. The uploaded image becomes part of the event card and detail page, helping users recognize the event quickly.",
        "QR code generation gave the ticket wallet a practical use. Even though the current verification is basic, the QR data creates a foundation for future scanning and entry management.",
    ],
    "2.7 Deployment Tools": [
        "Deployment planning was included so that the project does not remain only a local development folder. The root scripts install dependencies, build the client, and start the server in a way that suits a single-service deployment.",
        "The deployment files also helped me understand why environment variables are necessary. Secret values such as database URLs and JWT secrets should not be hardcoded in source code because they may change across local and production environments.",
    ],
    "CHAPTER 3 - INTRODUCTION TO PROJECT": [
        "EventfulEMS was selected as the main project because it has a clear real-world use case and enough technical depth for industrial training. It includes user management, forms, file upload, database relationships, API communication, conditional rendering, and a booking workflow.",
        "The project is also suitable for analytics because each user action creates structured information. Events, bookings, capacities, prices, and likes are not just screen elements; they are data points that can later be converted into reports.",
    ],
    "3.1 Project Overview": [
        "The project follows a straightforward user journey. A visitor first sees upcoming events. After login, the user can create an event or book an existing one. After booking, the generated ticket is stored in the wallet. This simple journey keeps the application easy to use.",
        "The application also supports the organizer side through the My Events page. This gives event creators a place to view their own events and delete them if required. This is important because a useful platform must support both attendees and organizers.",
    ],
    "3.2 Problem Statement": [
        "In many college-level events, information is spread across posters, chat groups, Google Forms, payment screenshots, and manual spreadsheets. This creates confusion for both organizers and participants. A participant may miss an update, and an organizer may struggle to track confirmed attendees.",
        "Another issue is that manual systems do not automatically preserve useful event data. Once the event is over, it becomes difficult to know how many people were interested, how many actually booked, what capacity was used, and which type of event performed better.",
    ],
    "3.3 Objectives of the Project": [
        "The project objectives were decided by looking at the complete event lifecycle. The system should support event creation, event discovery, ticket booking, ticket storage, and basic verification. Each objective was implemented as a separate module.",
        "A secondary objective was to keep the system understandable for future improvement. The project can later include analytics dashboards, payment gateway integration, admin roles, and QR scanning without changing the entire base structure.",
    ],
    "3.4 Proposed System": [
        "The proposed system stores all important event details in the database and displays them through a web interface. This reduces dependency on manual records and makes event information easier to update and access.",
        "The backend acts as the controller of important rules. For example, the frontend can request a ticket, but the backend decides whether the request is valid. This keeps important logic away from the browser and makes the system more reliable.",
    ],
    "3.5 Scope of the Project": [
        "The project scope is intentionally focused on essential features. It does not try to become a very large platform in the first version. Instead, it completes the basic event workflow properly and leaves space for future upgrades.",
        "The current scope is enough for demonstration in an academic training report because it shows frontend work, backend logic, database use, authentication, file upload, ticket generation, and analytics-ready data design.",
    ],
    "3.6 Project Modules": [
        "Dividing the project into modules made implementation manageable. When one feature had an issue, I could isolate it instead of searching across the entire application. This was especially helpful during ticket booking because that workflow touches frontend state, backend validation, event data, and ticket data.",
        "The modules also make the report easier to understand. A reader can study each part separately and then see how they connect to form the complete application.",
    ],
    "3.7 Authentication Module": [
        "The login and registration pages are important because almost every meaningful action depends on user identity. The system must know who created an event, who booked a ticket, and whose wallet should display which ticket.",
        "The UserContext on the frontend helps avoid repeated code. Once the profile is loaded, multiple pages can use the same user state. This made protected routes and user-specific actions easier to manage.",
    ],
    "3.8 Event Creation Module": [
        "The create event form is one of the most important organizer-facing screens. It collects all details needed for participants to decide whether they want to attend the event.",
        "The uploaded image improves event visibility. In event management, presentation matters because users are more likely to click an event when the card has a clear image, date, title, and price.",
    ],
    "3.9 Event Listing and Discovery": [
        "The listing page filters out past events so that users mainly see relevant upcoming events. This keeps the home page cleaner and prevents old events from distracting the user.",
        "The event card design also supports quick scanning. Users can compare event date, time, price, organizer, and popularity through likes before opening the detail page.",
    ],
    "3.10 Event Detail and Sharing": [
        "The event detail page is designed to provide enough information before booking. It separates the larger image, main title, price, description, organizer, date, time, and location so that users do not have to search for important details.",
        "Sharing options were included because events often grow through peer sharing. Copy link, WhatsApp share, and Facebook share actions help users pass the event to friends.",
    ],
    "3.11 Order Summary Module": [
        "The order summary page reduces mistakes by giving the user a final check before payment. It shows the ticket quantity, total amount, and event terms in one place.",
        "The maximum ticket quantity rule also supports fairness. A single user cannot book more than the defined limit, and the backend enforces this rule again during ticket creation.",
    ],
    "3.12 Payment Summary and Ticket Generation": [
        "The payment summary page completes the booking journey. Although real payment is not integrated yet, the screen represents the flow of confirming user details, checking the total, and creating the ticket.",
        "The ticket code is generated in the backend using time-based and random values. This makes every ticket easier to identify and avoids depending only on MongoDB's internal ID.",
    ],
    "3.13 Ticket Wallet Module": [
        "The ticket wallet works like a personal record page for attendees. It is useful because users may book more than one event and need a single place to view all tickets.",
        "Each wallet card contains both human-readable details and QR data. This makes the ticket useful for both the user and any future verification process.",
    ],
    "3.14 My Events Module": [
        "The My Events page helps organizers manage their own work. Without this page, a user would have to search through all public events to find the ones created by them.",
        "The delete confirmation protects event records from accidental removal. This is a small feature, but it improves trust because deletion is a permanent action.",
    ],
    "3.15 Calendar View Module": [
        "The calendar view gives a different way to explore events. Some users prefer list-based browsing, while others prefer checking dates visually. Both views use the same event data from the backend.",
        "This module also shows how one database collection can support multiple interface designs. The same event records are shown as cards on the home page and as date entries on the calendar page.",
    ],
    "3.16 Verification Center": [
        "The verification center is currently simple, but it is an important idea. Ticket systems are not complete if tickets can only be created but never checked. Even a basic verification page shows the direction of the project.",
        "In the future, this module can be improved by allowing authorized organizers to scan QR codes, mark tickets as used, and prevent repeated entry using the same ticket.",
    ],
    "3.17 Data Analytics Perspective": [
        "The analytics value of the project comes from its structured records. Every booking connects a user with an event and includes quantity, price, date, ticket code, and status. This is the type of data that can later be summarized in dashboards.",
        "For example, if an organizer wants to know which events perform best, the system can compare sold count, available quantity, likes, and revenue. If a college wants to plan future events, this data can help identify demand patterns.",
    ],
    "3.18 Important Data Entities": [
        "The User entity is simple, but it is central because it connects identity to actions. The Event entity stores public event information and organizer ownership. The Ticket entity stores the actual booking transaction.",
        "Together, these entities represent a complete data model for the current system. The model is also flexible enough to support additional fields such as event category, payment status, attendance status, feedback rating, and cancellation reason.",
    ],
    "3.19 Capacity and Booking Analytics": [
        "Capacity analytics helps answer a very practical question: how full is the event? This matters for venue planning, crowd management, promotion strategy, and future decision-making.",
        "If an event sells slowly, organizers may promote it more. If it sells out quickly, they may consider a larger venue or repeated session in the future. These decisions become easier when the application stores count and quantity properly.",
    ],
    "3.20 Revenue and Price Analytics": [
        "Revenue analytics is useful even in a student project because it shows how operational data can become business information. A ticket price is not only a display value; once multiplied by quantity, it becomes a measurable transaction value.",
        "The system can later calculate total revenue, average booking value, paid versus free event ratio, and month-wise income. These reports would help organizers understand financial performance without maintaining separate spreadsheets.",
    ],
    "3.21 User Engagement Analytics": [
        "Likes are a simple form of engagement data. They show interest even when a user does not immediately book a ticket. This can help identify events that attract attention but need better pricing, timing, or description to convert interest into bookings.",
        "In a future version, engagement can also include views, shares, wishlists, and reminder clicks. Combining these signals with bookings would make the analytics dashboard much more meaningful.",
    ],
    "3.22 Backend API Design": [
        "The API design follows the main user actions. This makes the backend easy to reason about. For example, event creation goes through createEvent, ticket booking goes through tickets, and user-specific tickets go through tickets/user.",
        "Clear API routes are useful during debugging. When a frontend page fails, the developer can directly check which route is involved and whether the problem is in request data, response data, or database logic.",
    ],
    "3.23 Database Design": [
        "The database design is intentionally practical. It stores enough information to complete the current workflow while keeping the schema understandable. Over-designing the database at this stage would make the project harder to explain and maintain.",
        "The ticket collection is especially important because it acts like a history of bookings. Even if an event changes later, the ticket still stores the details that were used during booking, such as event name, date, time, and price.",
    ],
    "3.24 Frontend Architecture": [
        "The frontend architecture is based on pages because the application has clear screens. Each page handles one major responsibility, such as listing events, creating events, showing event details, or displaying tickets.",
        "This structure made development easier because each file had a clear purpose. It also makes future improvements easier because a developer can open the relevant page file and work on that feature directly.",
    ],
    "3.25 Backend Architecture": [
        "The backend currently exists mainly in one server file, which is acceptable for a compact academic project. It keeps the full API flow visible in one place, making it easier to understand during report evaluation.",
        "For a production version, the backend can be divided into route files, controller files, middleware files, and model files. That would improve maintainability when the number of features grows.",
    ],
    "3.26 Implementation Challenges": [
        "One challenge was maintaining correct user-specific data. The application should not show another user's wallet or events. This required careful use of user IDs and protected frontend routes.",
        "Another challenge was keeping ticket count correct. If the backend only created tickets but did not update event count, capacity analytics would become wrong. Therefore, ticket creation and event count update were handled together.",
    ],
    "3.27 Solutions Applied": [
        "The solution was to keep important checks in the backend. Frontend controls improve user experience, but backend validation protects the actual data. This approach was used for ticket limits and capacity checks.",
        "The application also uses confirmation dialogs for delete actions, clear success or error messages, and fallback display areas when an event image is missing. These details make the application feel more stable.",
    ],
    "3.28 Testing Approach": [
        "Testing was done as a user journey rather than only checking isolated functions. I tested whether a person can register, log in, create an event, view it, book a ticket, and then find that ticket in the wallet.",
        "This workflow-based testing helped catch practical issues. A feature may work alone but fail when connected with another feature. For example, ticket booking depends on event details, user details, QR generation, and backend ticket creation.",
    ],
    "3.29 Benefits of the Project": [
        "The biggest benefit of EventfulEMS is that it organizes event-related work in one place. Users do not need separate forms, screenshots, and manual lists for every event.",
        "For learning, the project gave me experience with a complete application lifecycle. I understood how a feature starts as a UI idea, becomes an API request, is validated by the backend, stored in the database, and then displayed again to the user.",
    ],
    "3.30 Limitations of Current Version": [
        "The current version should be seen as a strong academic prototype rather than a finished commercial product. It demonstrates the main logic clearly, but some production features are still pending.",
        "The most important pending areas are real payments, advanced access control, persistent image storage, complete QR scanning, admin dashboards, and stronger reporting features.",
    ],
    "CHAPTER 4 - DESIGN, DIAGRAMS, AND SCREENSHOTS": [
        "This chapter should become more visual after final screenshots are inserted. The placeholders are intentionally included so that the final document can be completed neatly without changing the order of sections.",
        "While adding images, each screenshot should be cropped clearly and placed inside the placeholder area. Avoid adding screenshots that are too small, too dark, or filled with unnecessary browser tabs.",
    ],
    "CHAPTER 5 - TESTING AND RESULTS": [
        "Testing was kept simple but practical. The goal was to check the actual behavior of the system from the user's point of view. Since this is a web application, a successful result means the user can complete the workflow without confusion.",
        "The most important result is that the project is usable as a connected system. Events created in one module appear in other modules, and tickets created from bookings appear in the wallet.",
    ],
    "5.1 Functional Test Cases": [
        "The test cases cover the main workflow. They do not represent every possible edge case, but they verify the core features required for the project demonstration.",
        "More automated tests can be added in the future using frontend testing tools and API testing tools. For this version, manual testing was enough to validate the working prototype.",
    ],
    "5.2 Validation and Error Handling": [
        "Validation improves data quality. If missing or incorrect data is allowed into the database, later analytics will also become incorrect. Therefore, validation is not only a technical requirement but also an analytics requirement.",
        "Error messages also improve user experience. A user should know why an action failed and what can be done next instead of facing a silent failure.",
    ],
    "5.3 Results Achieved": [
        "The completed project shows that a MERN stack application can solve a real event management problem with a clean workflow. It also shows how full stack development and data analytics thinking can be combined.",
        "The project is ready for future improvements because the main data entities and workflows are already present. New dashboards, reports, and admin features can be built using the existing event and ticket data.",
    ],
    "CHAPTER 6 - FUTURE SCOPE AND CONCLUSION": [
        "The future scope is strong because event management systems can grow in many directions. The current project covers the base workflow, while future updates can focus on automation, analytics, security, and deployment quality.",
        "This project also gave me confidence to work on complete applications. It helped me understand that good software is not only about writing code, but also about designing a clear workflow and maintaining useful data.",
    ],
    "6.1 Future Scope": [
        "A future analytics dashboard can become the most valuable upgrade. It can show total events, total bookings, revenue, top events, upcoming events, sold-out events, and monthly activity.",
        "Role-based access can also make the platform more realistic. Admins can approve events, organizers can manage their own events, and attendees can only view and book tickets.",
    ],
    "6.2 Conclusion": [
        "Overall, the internship helped me improve technical ability, project discipline, and documentation skills. I learned how to convert an idea into a working system and how to explain that system in a formal report.",
        "EventfulEMS reflects my main contribution during the training period. It is simple enough to understand but complete enough to demonstrate full stack development, user workflows, backend logic, database design, and analytics potential.",
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="9CA3AF", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EventfulEMS Industrial Training Report")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p._p.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = p.add_run("EventfulEMS Industrial Training Report    Page ")
    label.font.name = "Times New Roman"
    label.font.size = Pt(9)
    label.font.color.rgb = RGBColor(90, 90, 90)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def para(doc, text="", bold=False, italic=False, align=None, size=12, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def placeholder(doc, title, instruction, height_lines=6):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell, "6B7280", "10")
    set_cell_shading(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(instruction)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    for _ in range(height_lines):
        cell.add_paragraph("")
    doc.add_paragraph()


def simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_cover(doc):
    for _ in range(2):
        para(doc)
    para(doc, "GURU NANAK DEV UNIVERSITY, AMRITSAR", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    para(doc, "Six Months Industrial Training Report", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    para(doc, "ON", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para(doc, "EVENTFULEMS", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, color="1F4E79")
    para(doc, "Event Management and Analytics System", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    para(doc)
    para(doc, "Project submitted for partial fulfillment for the degree of", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "BACHELOR OF TECHNOLOGY", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para(doc, "In Computer Science & Engineering", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para(doc, "Department of Computer Engineering and Technology", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    placeholder(doc, "COLLEGE / UNIVERSITY LOGO", "Insert official college logo here.", 3)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Submitted To:", "Submitted By:"),
        ("Project Mentor: __________________", "Ashika Kocher\nB.Tech CSE Sem 8\nRoll No.: 17032206387"),
    ]
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if r_idx == 0:
                run.bold = True
    para(doc, "Training Organization: Sabudh Foundation", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Training Duration: 8 January 2026 to Ongoing", align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def add_front_matter(doc):
    heading(doc, "CERTIFICATE", 1)
    para(doc, "This is to certify that the Industrial Training Report entitled \"EventfulEMS: Event Management and Analytics System\" is an authentic record of work carried out by Ashika Kocher during her six months of industrial training.")
    para(doc, "The training was undertaken at Sabudh Foundation from 8 January 2026 to ongoing, under the guidance and mentorship provided during the internship period. The work described in this report includes the learning experience, project development, implementation details, testing, and future scope of the project.")
    para(doc)
    para(doc, "Project Mentor Signature: __________________________")
    para(doc, "Date: __________________________")
    page_break(doc)

    heading(doc, "ONGOING INTERNSHIP CERTIFICATE", 1)
    para(doc, "This page is intentionally kept for attaching the ongoing internship certificate issued by Sabudh Foundation.")
    para(doc, "Instruction: Insert or paste the certificate image/PDF screenshot on this page after receiving the latest certificate from the organization. Keep the certificate centered and resize it so that the complete certificate, signature, and dates are visible clearly.")
    placeholder(doc, "CERTIFICATE PLACEHOLDER", "Attach ongoing internship certificate from Sabudh Foundation here.", 18)
    page_break(doc)

    heading(doc, "DECLARATION", 1)
    para(doc, "I hereby declare that the Industrial Training Report entitled \"EventfulEMS: Event Management and Analytics System\" is an authentic record of my own work completed as part of the six months industrial training requirement for the award of the degree of B.Tech. in Computer Science and Engineering.")
    para(doc, "The information presented in this report is based on my training experience at Sabudh Foundation and the project work developed during the internship. All references used for understanding the organization, technologies, and project concepts have been acknowledged in the reference section.")
    para(doc)
    para(doc, "Candidate Signature: __________________________")
    para(doc, "Name: Ashika Kocher")
    para(doc, "Roll No.: 17032206387")
    page_break(doc)

    heading(doc, "ACKNOWLEDGEMENT", 1)
    para(doc, "The successful completion of this industrial training report would not have been possible without the support, guidance, and encouragement of several people. I would like to express my sincere gratitude to my project mentor for giving me the opportunity to present my training work and for providing continuous guidance during the preparation of this report.")
    para(doc, "I am thankful to Sabudh Foundation for providing a learning environment where I could work on real-world technology concepts, improve my understanding of data science and analytics, and apply my full stack development skills to a practical project.")
    para(doc, "I would also like to thank my teachers, classmates, friends, and family for their support throughout the training period. Their encouragement helped me stay consistent while learning new tools, solving implementation issues, and improving the final project.")
    para(doc, "Ashika Kocher", True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    page_break(doc)

    heading(doc, "PREFACE", 1)
    para(doc, "Industrial training is an important part of engineering education because it connects classroom knowledge with real implementation. During this training, I worked on two areas: a short research-oriented robotics task titled Vision Language Navigation and the main project titled EventfulEMS.")
    para(doc, "EventfulEMS is the central focus of this report. It is an event management and analytics system built using the MERN stack. The project allows users to create events, view upcoming events, book tickets, manage personal tickets, and track useful event-related data such as ticket capacity, booking count, pricing, and event ownership.")
    para(doc, "The report is written in a simple and structured manner. Technical explanations are included only where they help understand the project, and separate placeholders have been provided wherever screenshots, diagrams, or output images should be added later.")
    page_break(doc)


def add_index(doc):
    heading(doc, "TABLE OF CONTENTS", 1)
    rows = [
        ("1", "Introduction", "1"),
        ("1.1", "About Sabudh Foundation", "1"),
        ("1.2", "Introduction to Internship", "2"),
        ("1.3", "Internship Learning Experience", "3"),
        ("1.4", "Vision Language Navigation Work", "4"),
        ("2", "Technology Used", "5"),
        ("2.1", "React and Vite", "5"),
        ("2.2", "Node.js and Express", "6"),
        ("2.3", "MongoDB and Mongoose", "7"),
        ("2.4", "Authentication, APIs, and Deployment", "8"),
        ("3", "Introduction to Project", "10"),
        ("3.1", "Project Overview", "10"),
        ("3.2", "Problem Statement and Objectives", "11"),
        ("3.3", "Modules of EventfulEMS", "13"),
        ("3.4", "Data Analytics Perspective", "19"),
        ("3.5", "Implementation Details", "24"),
        ("4", "Design, Diagrams, and Screenshots", "33"),
        ("5", "Testing and Results", "42"),
        ("6", "Limitations, Future Scope, and Conclusion", "46"),
        ("7", "References", "50"),
    ]
    simple_table(doc, ["Sr. No.", "Contents", "Page No."], rows, [Cm(2), Cm(12), Cm(3)])
    page_break(doc)


def add_page(doc, title, paragraphs=None, bullets=None, table=None, image=None, level=1, break_after=False):
    heading(doc, title, level)
    for text in paragraphs or []:
        para(doc, text)
    for text in bullets or []:
        bullet(doc, text)
    if table:
        simple_table(doc, table[0], table[1], table[2] if len(table) > 2 else None)
    if image:
        placeholder(doc, image[0], image[1], max(image[2] if len(image) > 2 else 6, 24))
    if break_after:
        page_break(doc)


def build_report():
    doc = Document()
    set_doc_defaults(doc)
    add_footer(doc.sections[0])
    add_cover(doc)

    # Restart document section for the body/front matter footer.
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin = Cm(2.0)
    body_section.bottom_margin = Cm(2.0)
    body_section.left_margin = Cm(2.5)
    body_section.right_margin = Cm(2.5)
    add_page_number_footer(body_section)

    add_front_matter(doc)
    add_index(doc)

    def spec(title, paragraphs, **kwargs):
        return (title, paragraphs, kwargs)

    pages = [
        spec("CHAPTER 1 - INTRODUCTION", [
            "Industrial training helped me connect academic computer science concepts with real project development. During this period, I worked as a Data Science intern at Sabudh Foundation from 8 January 2026 to ongoing.",
            "The main project developed during this training is EventfulEMS, an event management and analytics system. It is designed for institutions, clubs, and organizers who need a simple way to publish events, manage bookings, and understand event participation through structured data.",
            "This report explains the organization, internship learning, technologies used, project design, implementation, testing, limitations, and future scope."
        ]),
        spec("1.1 About Sabudh Foundation", [
            "Sabudh Foundation is a research-driven skilling initiative that focuses on ethical AI, data science education, and project-based learning. Its programs are designed around guided mentorship, hands-on work, and practical problem solving.",
            "The foundation offers six-month programs in areas such as Data Science, Data Analytics, and AIoT. The official website describes its approach as learning through real-world projects, live mentoring, and community-based innovation.",
            "For me, Sabudh provided a structured environment to improve both technical skills and problem-solving confidence. The internship encouraged consistent practice, independent research, and clear communication of project work."
        ]),
        spec("1.2 Introduction to Internship", [
            "The internship was taken in the Data Science track at Sabudh Foundation. The duration of the training was planned as six months, starting from 8 January 2026 and continuing through the current semester.",
            "The work included learning sessions, technical practice, project discussions, and implementation tasks. Although the main project is a full stack application, it was treated from a data analytics viewpoint because the system collects event, booking, ticket, capacity, and user activity data.",
            "The training helped me understand how a software system can be developed not only for user interaction, but also for collecting clean operational data that can later be used for reports and decision-making."
        ]),
        spec("1.3 Internship Learning Experience", [
            "During the internship, I strengthened my understanding of web development, backend APIs, database design, authentication, deployment, and analytics-oriented thinking. I also learned how to keep a project practical instead of overcomplicating it.",
            "The project required breaking a larger idea into smaller modules such as user login, event creation, event listing, ticket booking, QR ticket generation, calendar view, and verification. Working module by module made the project easier to debug and improve.",
            "A major learning outcome was understanding that a useful application should be simple for users, but carefully structured behind the scenes. Data models, API routes, validation, and error handling all matter because they directly affect reliability."
        ]),
        spec("1.4 Vision Language Navigation Work", [
            "Apart from EventfulEMS, I also explored a robotics-related project titled Vision Language Navigation. The aim of VLN is to build an AI agent that can understand human instructions and navigate inside indoor spaces.",
            "The concept combines language understanding, visual perception, and navigation planning. For example, an instruction such as \"go up the stairs and stop near the mirror\" requires the system to identify objects, understand spatial cues, decide a path, and move step by step.",
            "The project used concepts from the Matterport3D Room-to-Room dataset, which includes real indoor scans, navigation graphs, RGB observations, human-written instructions, and ground truth paths. My work mainly involved understanding the flow of perception, reasoning, path planning, and evaluation metrics such as success rate and SPL."
        ]),
        spec("CHAPTER 2 - TECHNOLOGY USED", [
            "EventfulEMS was developed using the MERN stack. The frontend was built with React and Vite, the backend with Node.js and Express, and the database layer with MongoDB and Mongoose.",
            "The technology choices were practical for a student-level industrial project because they support fast development, API-based communication, reusable UI components, and flexible document-based data storage.",
            "This chapter explains the major tools used in the project and how each one contributed to the final system."
        ]),
        spec("2.1 React and Vite", [
            "React was used to build the frontend user interface. It allowed the project to be divided into pages and reusable components such as the header, footer, confirmation dialog, event cards, ticket wallet, and forms.",
            "Vite was used as the development build tool. It provides a faster local development experience and simple production build support. This made frontend testing and UI iteration smoother during project development.",
            "React Router was used to create routes for the home page, login, registration, event details, create event page, calendar view, ticket wallet, order summary, payment summary, and verification center."
        ]),
        spec("2.2 Tailwind CSS and Frontend Styling", [
            "Tailwind CSS was used for styling the application. It helped create responsive layouts, buttons, event cards, forms, confirmation dialogs, and dashboard-like views without writing large separate CSS files.",
            "The frontend design focuses on clarity. Events are displayed as cards with image, title, date, time, organizer, price, and booking link. The calendar view gives users another way to scan events by date.",
            "Icons from React Icons were used for calendar, location, delete, sharing, and navigation actions. These visual cues made the interface easier to understand."
        ]),
        spec("2.3 Node.js and Express", [
            "Node.js was used for the backend runtime and Express was used to define API routes. The backend handles user registration, login, logout, profile fetching, event creation, event listing, ticket booking, ticket deletion, and event deletion.",
            "Express made it easier to structure the server as a set of API endpoints. The project uses a router mounted under /api for production and also supports direct routes in development.",
            "The backend also serves uploaded event images from the uploads folder and serves the frontend build in production when deployed as a single service."
        ]),
        spec("2.4 MongoDB and Mongoose", [
            "MongoDB was used as the database because EventfulEMS stores flexible records such as users, events, and tickets. Mongoose was used to define schemas and interact with the database in a structured way.",
            "The User schema stores name, email, and hashed password. The Ticket schema stores user ID, event ID, ticket code, booking details, QR code data, quantity, status, and timestamps.",
            "The event schema stores details such as title, description, organizer, event date, time, location, ticket price, ticket quantity, image, owner information, likes, and booking count."
        ]),
        spec("2.5 Authentication and Security Tools", [
            "JWT authentication was used to keep users signed in. After registration or login, the backend signs a token and stores it in an HTTP-only cookie. This helps protect the token from direct frontend access.",
            "Passwords are hashed using bcryptjs before being stored in MongoDB. This is a basic but important security practice because plain text passwords should never be saved in a database.",
            "CORS and cookie settings were configured so that the frontend and backend can communicate during local development and production deployment."
        ]),
        spec("2.6 API, Upload, and QR Tools", [
            "Axios was used on the frontend to communicate with backend API routes. It is used for fetching events, creating events, logging in, booking tickets, deleting tickets, and loading user-specific data.",
            "Multer was used in the backend to handle event image uploads. Uploaded images are saved in the uploads directory and then displayed in event cards and event detail pages.",
            "The qrcode package was used to create QR code data for booked tickets. This gave each ticket a scannable visual identity and made the ticket wallet more practical."
        ]),
        spec("2.7 Deployment Tools", [
            "The project includes deployment support for Render. The root package file contains scripts for installing backend and frontend dependencies, building the client, and starting the API server.",
            "The Express server is configured to serve the React production build in production mode. This allows the app to run as one Render web service instead of needing separate deployments for frontend and backend.",
            "Environment variables such as MONGO_URL, JWT_SECRET, CLIENT_ORIGIN, PORT, and MONGO_DB_NAME are used to keep deployment configuration separate from source code."
        ]),
        spec("CHAPTER 3 - INTRODUCTION TO PROJECT", [
            "EventfulEMS is an event management and analytics system designed to simplify how events are created, discovered, booked, and tracked. The project was developed as the main work of my industrial training.",
            "The system supports two major user perspectives. First, attendees can browse events, view details, book tickets, and manage their ticket wallet. Second, organizers can create events, track capacity, and manage their own event listings.",
            "From an analytics perspective, the project captures data that can answer useful questions such as which events are active, how many tickets are sold, how much revenue is generated, and which events have higher user engagement."
        ]),
        spec("3.1 Project Overview", [
            "The application begins with a home page that displays upcoming events. Each event card contains key information such as title, date, time, price, organizer, creator, image, likes, and capacity progress.",
            "Users can open an event detail page, read the event description, check the date and venue, share the event link, and proceed to book tickets. The order summary page ensures that the user accepts event terms before moving forward.",
            "After checkout, the system creates a ticket with booking details, total price, QR code, and a unique ticket code. The ticket then appears in the user wallet."
        ]),
        spec("3.2 Problem Statement", [
            "Many campus and community events are still managed using scattered forms, messages, manual lists, and separate payment or ticket records. This makes it difficult to maintain a clean event history or understand participation trends.",
            "Organizers need a system where event details, ticket capacity, bookings, and user records stay connected. Attendees also need a simple way to discover events and keep their tickets in one place.",
            "The problem addressed by EventfulEMS is to create a simple digital platform that combines event publishing, booking, ticket storage, and basic operational analytics in one application."
        ]),
        spec("3.3 Objectives of the Project", [
            "The main objective of EventfulEMS is to provide a complete event workflow from event creation to ticket booking and verification.",
            "The project also aims to make event data more organized so that future analytics can be performed on booking count, ticket capacity, event categories, pricing, and user participation.",
            "The system was designed to be understandable, deployable, and extendable. It avoids unnecessary complexity while still covering the major functions expected from an event management platform."
        ], bullets=[
            "Allow authenticated users to create events with images and ticket information.",
            "Allow attendees to browse, like, share, and book upcoming events.",
            "Generate QR-based ticket data and unique ticket codes.",
            "Provide a wallet where users can view and delete their booked tickets.",
            "Support basic capacity tracking through ticket count and total quantity."
        ]),
        spec("3.4 Proposed System", [
            "The proposed system is a web-based event management application built using React, Node.js, Express, and MongoDB. It connects user actions with database records through REST APIs.",
            "When an organizer creates an event, the event record is saved in MongoDB and becomes visible on the home page and calendar page. When a user books a ticket, the system validates the event, checks the ticket limit, updates the sold count, and stores the ticket.",
            "The proposed system improves manual event workflows by keeping event information, booking records, ticket codes, and capacity data in one place."
        ]),
        spec("3.5 Scope of the Project", [
            "The scope of EventfulEMS includes frontend design, backend API development, MongoDB database integration, authentication, ticket booking, ticket wallet, event management, and deployment preparation.",
            "The current version is suitable for academic, club, and small organizational event workflows. It can be used as a foundation for a more advanced analytics platform in the future.",
            "Advanced features such as real payment gateway integration, admin analytics dashboard, email ticket delivery, and role-based organizer approvals are part of the future scope."
        ]),
        spec("3.6 Project Modules", [
            "The project was divided into smaller modules so that each feature could be implemented and tested independently. This modular approach helped reduce confusion during development.",
            "The main modules include authentication, event listing, event creation, event detail view, order summary, payment summary, ticket wallet, calendar view, my events page, and verification center.",
            "Each module contributes to the overall workflow and also creates or uses structured data that supports future analytics."
        ], table=(["Module", "Purpose"], [
            ("Authentication", "Register users, log in, log out, and load profile data."),
            ("Event Creation", "Create event records with title, date, location, image, price, and quantity."),
            ("Ticket Booking", "Create tickets, calculate price, and update event sold count."),
            ("Ticket Wallet", "Display booked tickets with QR code and ticket ID."),
            ("Calendar View", "Show events by month and date for easy planning."),
        ], [Cm(5), Cm(11)])),
        spec("3.7 Authentication Module", [
            "The authentication module allows users to register and log in. When a user registers, the password is hashed before being stored. On successful login, the backend creates a JWT token and stores it in a cookie.",
            "The frontend uses the user context to know whether the user is logged in. Protected pages such as create event, my events, ticket wallet, and verification center redirect the user to login when needed.",
            "This module is important because bookings and event ownership must be linked to specific users."
        ]),
        spec("3.8 Event Creation Module", [
            "The event creation module allows a logged-in user to publish a new event. The form collects event title, optional information, description, organizer name, date, time, location, ticket price, total tickets, and image.",
            "The form data is sent to the backend as multipart form data because it includes an image file. The backend stores the image and saves the event record in MongoDB.",
            "The system also stores owner and owner name values so that each event can be connected to the user who created it."
        ]),
        spec("3.9 Event Listing and Discovery", [
            "The home page fetches event records from the backend and displays upcoming events. Events that have already passed are not shown in the main upcoming event grid.",
            "Each card includes useful summary details: event image, title, date, time, price, description, organizer, creator, likes, and a booking link. The capacity bar gives a quick visual indication of tickets sold compared to total quantity.",
            "This page is the main entry point for attendees and helps them scan available events quickly."
        ]),
        spec("3.10 Event Detail and Sharing", [
            "The event detail page gives a focused view of one event. It displays the uploaded image, title, ticket price, description, organizer, date, time, and location.",
            "The page also includes sharing options such as copying the event link and opening WhatsApp or Facebook share actions. These features are useful because event discovery often depends on social sharing.",
            "From this page, the user can proceed to the order summary before booking tickets."
        ]),
        spec("3.11 Order Summary Module", [
            "The order summary module acts as a confirmation step before payment. It displays the event name, price, selected ticket quantity, available ticket count, subtotal, and important terms and conditions.",
            "The user can select between 1 and 10 tickets, subject to availability. This protects the system from accidental large bookings and keeps ticket distribution fair.",
            "The proceed button stays disabled until the user accepts the terms and conditions. This small workflow step improves user clarity before ticket creation."
        ]),
        spec("3.12 Payment Summary and Ticket Generation", [
            "The payment summary page collects user and payment-style details and then creates a ticket record. In the current project, the payment interface is a simulated checkout screen used to complete the booking workflow.",
            "Before saving a ticket, the frontend generates QR code data using the qrcode package. The backend then validates the request, checks ticket limits, calculates total price, creates a unique ticket code, saves the ticket, and updates the event sold count.",
            "After successful booking, the user sees a confirmation screen with the ticket code and options to view the wallet or browse more events."
        ]),
        spec("3.13 Ticket Wallet Module", [
            "The ticket wallet displays all tickets booked by the logged-in user. Each ticket card shows QR code, event name, date, time, user name, quantity, price, email, and ticket ID.",
            "The wallet gives users a single place to manage booked tickets. Users can also delete a ticket after confirming through a dialog box.",
            "This module is important from the analytics point of view because it shows how ticket records connect user data, event data, quantity, and pricing."
        ]),
        spec("3.14 My Events Module", [
            "The My Events page displays events created by the logged-in user. It allows organizers to view their own event listings, see date and location information, check ticket count and capacity, and delete events when required.",
            "This module supports the organizer side of the application. It gives creators a direct way to manage the events they have published.",
            "The page also uses a confirmation dialog before deletion so that users do not accidentally remove event records."
        ]),
        spec("3.15 Calendar View Module", [
            "The calendar view displays events month-wise. Users can move to the previous or next month and see event titles placed on their event dates.",
            "This module is useful for planning because users often think about events by date rather than by list order. It also makes the application feel more complete as an event management tool.",
            "The calendar uses date-fns for month calculations, day formatting, and event date matching."
        ]),
        spec("3.16 Verification Center", [
            "The verification center allows a logged-in user to enter a ticket ID and check whether it matches a ticket in their wallet.",
            "In the present version, verification is user-specific and simple. It checks the entered ticket code against the user's ticket records and displays whether a matching ticket exists.",
            "This module can later be expanded into an organizer-facing scanning or gate-entry feature where authorized staff can verify any event ticket."
        ]),
        spec("3.17 Data Analytics Perspective", [
            "Although EventfulEMS is implemented as a full stack system, it has a strong data analytics angle. The application stores structured event and ticket data that can be analyzed to understand event performance.",
            "Important measurable fields include event date, location, organizer, ticket price, total quantity, sold count, likes, user ID, ticket quantity, total price, and ticket status.",
            "With these fields, the system can later generate dashboards showing popular events, revenue trends, booking patterns, sold-out risk, and engagement."
        ]),
        spec("3.18 Important Data Entities", [
            "The three main data entities are User, Event, and Ticket. These entities are connected through user ID and event ID fields.",
            "User data identifies the person using the system. Event data describes what is being organized. Ticket data records the booking transaction and connects a user to an event.",
            "This structure is simple but effective because it supports the complete workflow and can be extended for analytics without redesigning the entire system."
        ], table=(["Entity", "Key Fields", "Analytics Use"], [
            ("User", "name, email", "Understand registered users and ownership."),
            ("Event", "title, date, location, price, quantity, count, likes", "Measure event demand and capacity usage."),
            ("Ticket", "userid, eventid, ticketCode, count, totalPrice, status", "Track bookings, revenue, and attendance flow."),
        ], [Cm(3), Cm(7), Cm(6)])),
        spec("3.19 Capacity and Booking Analytics", [
            "Capacity tracking is one of the most useful analytics features in EventfulEMS. Each event stores total ticket quantity and the current sold count.",
            "Whenever a ticket is booked, the backend checks whether enough tickets are available and then increases the sold count. The frontend displays this as a progress bar on event cards.",
            "This helps organizers understand how close an event is to being full and can support future decisions such as increasing venue size, closing registrations, or promoting low-booking events."
        ]),
        spec("3.20 Revenue and Price Analytics", [
            "Ticket price and total price fields make it possible to calculate event revenue. For free events, the system records price as zero, while paid events can calculate total based on ticket quantity.",
            "The ticket record stores the final total price so that revenue can be studied later without recalculating everything from the event page.",
            "Future dashboards can use this data to show total revenue per event, average ticket price, paid versus free event performance, and monthly booking value."
        ]),
        spec("3.21 User Engagement Analytics", [
            "The system stores likes for events and connects bookings to users. Likes are a simple engagement signal that show which events attract interest even before booking.",
            "Booking count, likes, and event views can later be combined to understand which events have high visibility but low conversion, or which events sell tickets quickly.",
            "These analytics would help organizers improve event descriptions, pricing, images, and promotion strategy."
        ]),
        spec("3.22 Backend API Design", [
            "The backend API was designed around the main actions of the system. There are routes for registration, login, profile, logout, event creation, event fetching, event deletion, ticket creation, ticket fetching, and ticket deletion.",
            "The API accepts JSON for most requests and multipart form data for event image upload. It returns JSON responses so the frontend can update the interface without reloading the page.",
            "The route structure is simple and readable, which made debugging easier during development."
        ], table=(["API Route", "Method", "Purpose"], [
            ("/api/register", "POST", "Create user and sign token."),
            ("/api/login", "POST", "Authenticate user."),
            ("/api/createEvent", "POST/GET", "Create or fetch events."),
            ("/api/event/:id", "GET/DELETE", "Fetch or delete one event."),
            ("/api/tickets", "POST", "Book ticket."),
            ("/api/tickets/user/:userId", "GET", "Fetch user's tickets."),
        ], [Cm(5), Cm(3), Cm(8)])),
        spec("3.23 Database Design", [
            "MongoDB stores the project data in collections. The flexible document model suits this project because event and ticket records include nested details, optional images, and changing fields.",
            "Mongoose schemas provide structure over MongoDB documents. This reduces mistakes by defining expected fields and validation rules for ticket data.",
            "The ticket schema is the most detailed schema because it stores booking details, QR code data, ticket code, quantity, status, and timestamps."
        ]),
        spec("3.24 Frontend Architecture", [
            "The frontend is organized into route-based pages. App.jsx defines the major routes, while individual page files handle event display, creation, booking, calendar view, wallet, and verification.",
            "A UserContext file manages user state across the application. This prevents repeated profile checks in every component and helps protected pages react to login status.",
            "Utility files are used for API and media URL handling. This keeps environment-specific URL logic away from the page components."
        ]),
        spec("3.25 Backend Architecture", [
            "The backend server starts by loading environment variables, setting up Express middleware, configuring CORS, serving uploads, validating required secrets, and connecting to MongoDB.",
            "Routes are grouped under an Express router and then mounted under /api. The same backend also supports serving the built React app in production mode.",
            "This architecture is compact and suitable for the project size. It can later be refactored into separate route and controller files if the project grows."
        ]),
        spec("3.26 Implementation Challenges", [
            "One implementation challenge was keeping frontend and backend paths consistent between local development and production. This was handled by using API utility functions and environment variables.",
            "Another challenge was connecting ticket booking with event capacity. The backend needed to check previous tickets, ticket limit, total quantity, and sold count before creating a new ticket.",
            "Image uploads also required attention because uploaded files must be saved by the backend and served correctly to the frontend."
        ]),
        spec("3.27 Solutions Applied", [
            "API base URL handling was centralized so that the frontend could work in local development and production. Cookie credentials were enabled in Axios to support authentication.",
            "Ticket creation was moved through backend validation. The backend checks user ID, event ID, per-user ticket limit, total event capacity, ticket code generation, and event count update.",
            "A static uploads route was added so event images can be accessed consistently from event cards and event detail pages."
        ]),
        spec("3.28 Testing Approach", [
            "Testing was performed module by module. First, the basic pages and navigation were checked. Then forms, API calls, authentication, event creation, ticket booking, and wallet display were tested.",
            "Manual testing was important because this project includes many user workflows. Each workflow was checked from the perspective of a normal user, such as creating an event, booking a ticket, and finding it in the wallet.",
            "Error cases were also considered, including unauthenticated access, missing event data, ticket limit reached, and unavailable tickets."
        ]),
        spec("3.29 Benefits of the Project", [
            "EventfulEMS provides a practical digital solution for managing events and tickets. It reduces scattered manual work and keeps event-related records together.",
            "The system is useful for students and organizers because it includes the full workflow: publish event, view event, book ticket, generate QR ticket, and manage records.",
            "It also creates a foundation for analytics by storing structured data about users, events, tickets, prices, quantities, and engagement."
        ]),
        spec("3.30 Limitations of Current Version", [
            "The current payment screen is a simulated workflow and does not connect to a real payment gateway. This is acceptable for the project demonstration but would need improvement for real use.",
            "The verification center is simple and user-specific. A production system would need organizer roles, scanner-based QR verification, and stronger ticket status handling.",
            "Uploaded images are stored locally. On free deployment platforms, local uploads may disappear after redeployment unless persistent storage is configured."
        ]),
        spec("3.31 Software Development Life Cycle", [
            "The development of EventfulEMS followed a practical and iterative software development life cycle. The project was not built in one step; it was divided into requirement understanding, planning, design, implementation, testing, and improvement.",
            "During requirement understanding, I identified the main users of the system: attendees, event organizers, and future administrators. Attendees need to discover and book events. Organizers need to publish events and track bookings. Administrators may later need approval and reporting features.",
            "In the planning stage, the project was broken into independent modules. Authentication was planned first because many other features depend on the active user. Event creation and event listing were planned next, followed by booking, wallet, calendar, and verification features.",
            "The design stage included deciding the frontend routes, backend API endpoints, database entities, and data flow. The goal was to keep the application understandable while still covering the complete event workflow.",
            "Implementation was done module by module. After completing each module, I tested it manually before moving to the next one. This reduced the chance of carrying unresolved bugs into later modules.",
            "The testing and improvement stage involved checking real user flows such as logging in, posting an event, booking a ticket, checking ticket limits, viewing the wallet, and deleting records. The project was improved whenever a missing validation or unclear user response was found."
        ]),
        spec("3.32 Feasibility Study", [
            "A feasibility study was important because a project should be realistic within the available time, tools, and skill level. EventfulEMS was technically feasible because the MERN stack provides all required features for frontend, backend, database, authentication, and deployment.",
            "The project was economically feasible because all main technologies used were open-source or free for development. React, Vite, Express, MongoDB, Mongoose, Tailwind CSS, and Render deployment support can be used without high setup cost for a student project.",
            "The project was operationally feasible because the workflow is familiar to users. Most students already understand event cards, booking buttons, ticket confirmation screens, and wallet-style pages. This reduced the learning curve for the application.",
            "The project was also scalable at a basic level. The current code can support more events and users as long as the database and hosting service are configured properly. For larger scale, image storage, payment gateway, caching, and role-based dashboards would need to be added.",
            "From an analytics point of view, the project is feasible because the required data is already collected during normal use. Booking analytics does not require extra manual data entry; it can be calculated from event and ticket records."
        ]),
        spec("3.33 Requirement Analysis", [
            "Requirement analysis helped convert the project idea into clear features. Functional requirements define what the system should do, while non-functional requirements define how well the system should perform.",
            "The main functional requirements were user registration, user login, event creation, event listing, event details, ticket booking, ticket wallet, ticket deletion, event deletion, calendar display, and ticket verification.",
            "Important non-functional requirements included usability, responsiveness, basic security, maintainability, and deployment readiness. The interface needed to be simple enough for students and organizers to use without training.",
            "Security requirements included password hashing, session handling, protected pages, and backend validation. Even though the project is academic, these practices make it closer to a real-world application.",
            "Data requirements included storing event details, user details, ticket details, ticket quantity, total ticket price, ticket code, QR code data, event capacity, sold count, and timestamps. These values support both application workflow and future analytics."
        ]),
        spec("3.34 Functional Requirements", [
            "The system should allow a new user to create an account using name, email, and password. The password should not be stored directly and must be hashed before saving.",
            "The system should allow a registered user to log in and access protected features. If a user is not logged in, protected pages should redirect to the login page.",
            "The system should allow logged-in users to create events with title, description, organizer, date, time, location, ticket price, total ticket quantity, and image.",
            "The system should show upcoming events on the home page and should allow users to open detailed event pages. The event detail page should include booking and sharing options.",
            "The system should allow users to choose ticket quantity within the allowed limit and book tickets only when seats are available.",
            "The system should generate ticket records with QR code, unique ticket code, user details, event details, quantity, price, and status.",
            "The system should allow users to view their own booked tickets in a wallet and delete tickets after confirmation.",
            "The system should allow organizers to view and delete their own created events."
        ]),
        spec("3.35 Non-Functional Requirements", [
            "Usability was a major non-functional requirement. The pages should be understandable and should guide the user naturally from browsing to booking.",
            "Responsiveness was also important because users may open the application on different screen sizes. Tailwind CSS was used to create layouts that adjust for different devices.",
            "Reliability was considered while designing ticket booking. The backend checks ticket quantity, event availability, and user ticket limits before creating a ticket record.",
            "Maintainability was supported by dividing the frontend into page components and utility files. Backend code can later be split into separate route and controller files as the project grows.",
            "Security was considered through hashed passwords, JWT cookies, environment variables, and protected routes. These features reduce common risks in a basic web application.",
            "Deployability was considered by adding Render configuration, production build scripts, and environment-based API handling."
        ]),
        spec("3.36 User Roles and Access", [
            "In the current version, the main user role is a registered user. A registered user can act as both attendee and organizer. This keeps the project simple while still supporting the full workflow.",
            "As an attendee, the user can browse events, open event pages, book tickets, view wallet tickets, and delete personal tickets. These actions are connected to the user's ID.",
            "As an organizer, the same user can create events and view the events created by them. The My Events page filters event records according to ownership fields.",
            "A future admin role can be added to approve events, remove inappropriate content, view platform-wide analytics, manage users, and handle event disputes.",
            "Role-based access control would make the project stronger because it would separate permissions clearly. Attendees should not access organizer tools unless they create events, and admins should have controlled platform-level access."
        ]),
        spec("3.37 Detailed Event Workflow", [
            "The event workflow begins when a logged-in user opens the create event page. The system fills owner information using the current user context so that the event can be linked to the creator.",
            "The user enters event details and uploads an image. The frontend sends this data as FormData because normal JSON cannot carry the image file directly.",
            "The backend receives the request through the create event API route. Multer stores the uploaded image in the uploads directory and the event details are saved in MongoDB.",
            "Once saved, the event becomes available to other pages. It appears on the home page, event detail page, calendar view, and my events page depending on the route being used.",
            "When a user opens the event detail page, the frontend fetches the event by ID. This keeps the displayed details synchronized with the database record.",
            "The workflow shows how one event record moves through the system and supports multiple user interfaces."
        ]),
        spec("3.38 Detailed Ticket Booking Workflow", [
            "The ticket booking workflow starts from the event detail page. The user clicks Book Ticket and is redirected to the order summary page.",
            "On the order summary page, the system displays event terms, ticket price, available quantity, selected quantity, and subtotal. The user must accept the confirmation checkbox before proceeding.",
            "The payment summary page collects user details and simulates payment information. Before creating the ticket, the frontend generates QR code data using the event and user information.",
            "The backend ticket route performs the most important validation. It checks whether user ID and event ID exist, whether the event exists, how many tickets the user already bought, and whether enough tickets are available.",
            "If the request is valid, the backend generates a unique ticket code and creates a ticket document. The event sold count is then updated so future users see the correct remaining capacity.",
            "This workflow connects frontend state, backend validation, MongoDB ticket storage, event count update, QR generation, and final wallet display."
        ]),
        spec("3.39 Analytics Dashboard Design Proposal", [
            "Although the current version does not include a complete dashboard, the data model already supports one. A dashboard can be added for organizers and administrators in future versions.",
            "The organizer dashboard can show total events created, total tickets sold, remaining tickets, total revenue, average ticket price, and top performing events.",
            "A chart can show bookings over time. This would help organizers understand whether registrations happen early, near the event date, or only after promotion.",
            "Another useful chart can compare free and paid events. This can show whether ticket price affects booking behaviour in a college environment.",
            "Capacity analytics can display occupancy percentage for each event. Events above 80 percent can be marked as high demand, while events below 30 percent can be marked for promotion.",
            "The dashboard can also show engagement analytics using likes and bookings. Events with many likes but fewer bookings may need better pricing or clearer information."
        ]),
        spec("3.40 Sample Analytics Metrics", [
            "Several metrics can be calculated from the current database fields. These metrics can convert raw data into meaningful event insights.",
            "Total Bookings can be calculated by summing ticket count values. This tells how many total seats have been booked across one event or across the platform.",
            "Revenue can be calculated by summing totalPrice values in ticket records. This can be grouped by event, month, organizer, or category in future versions.",
            "Occupancy Percentage can be calculated as sold count divided by total quantity multiplied by 100. This is useful for understanding event demand.",
            "Average Ticket Quantity can be calculated by dividing total booked tickets by number of ticket records. This shows whether users usually book single tickets or group tickets.",
            "Engagement Rate can be estimated by comparing likes with bookings. Although simple, this can help identify whether interest is converting into actual participation."
        ]),
        spec("3.41 Deployment and Environment Configuration", [
            "Deployment planning was included so that the project can run outside the local development machine. The project contains a Render configuration file and root-level build scripts.",
            "The backend reads important values from environment variables. MONGO_URL stores the database connection string, JWT_SECRET stores the token signing secret, PORT defines the backend port, and CLIENT_ORIGIN controls allowed frontend origins.",
            "The frontend reads the API base URL through Vite environment configuration. This allows the same code to work in local development and production.",
            "In production, the backend can serve the built React application from the client dist folder. This simplifies hosting because one web service can handle both frontend pages and backend APIs.",
            "The deployment documentation also notes an important limitation: uploaded event images stored on a free Render filesystem may disappear after redeployment. A production version should use persistent storage such as Cloudinary or S3."
        ]),
        spec("3.42 Maintenance and Upgrade Plan", [
            "Maintenance is an important part of any real project. After deployment, the application would need regular updates, bug fixes, dependency checks, database monitoring, and user feedback review.",
            "One maintenance task would be checking outdated npm packages and updating them carefully. Since frontend and backend both use JavaScript packages, dependency management is important.",
            "Another maintenance task would be monitoring failed bookings or API errors. Logs can help identify whether users are facing problems with authentication, ticket limits, image uploads, or database connectivity.",
            "Database maintenance would include indexing frequently searched fields such as event ID, user ID, ticket code, and event date. This would improve performance as data grows.",
            "Upgrade planning can include adding event editing, admin approvals, email notifications, QR scanning, analytics dashboards, and real payment gateway integration."
        ]),
        spec("3.43 Lessons Learned from Project Development", [
            "The project taught me that full stack development requires careful coordination between frontend, backend, and database. A small mismatch in field names can break a complete feature.",
            "I learned that user experience depends heavily on small details such as loading states, success messages, confirmation dialogs, redirects, and clear error messages.",
            "I also learned that backend validation is more reliable than frontend validation alone. The frontend can guide the user, but the backend must protect the database.",
            "The project improved my understanding of how data analytics begins at the application design stage. If the right fields are stored from the beginning, analytics becomes easier later.",
            "Finally, I learned the importance of writing clean documentation. A project report should make the project understandable to someone who has not seen the code."
        ]),
        spec("3.44 Data Cleaning and Reporting Considerations", [
            "For a data analytics project, raw data is only useful when it is consistent. EventfulEMS stores operational data through user actions, so the quality of analytics depends on how cleanly that data is collected.",
            "Dates should be stored in a consistent format so that events can be grouped by day, month, or year. The project stores event dates as date values, which makes calendar display and future date-based reporting easier.",
            "Numeric fields such as ticket price, total ticket quantity, sold count, and ticket count should always be stored as numbers. If these values are stored as text, calculations such as revenue and occupancy percentage become unreliable.",
            "User names, event names, and locations may need text cleaning in a future analytics layer. For example, extra spaces, different capitalization, or spelling variations can affect grouping in reports.",
            "A future reporting module can include simple data cleaning rules before generating charts. These rules can standardize event titles, normalize dates, remove invalid records, and handle missing values.",
            "This consideration is important because analytics is not only about creating charts. It also requires trust in the data that feeds those charts."
        ]),
        spec("3.45 Suggested Database Queries for Analytics", [
            "The current MongoDB collections can support several useful queries. These queries can be used later to build dashboards and reports.",
            "To find popular events, the system can sort events by ticket sold count or likes. This would show which events are receiving the most attention from users.",
            "To calculate revenue, the system can aggregate ticket records by event ID and sum totalPrice. This would show total earning per event.",
            "To find occupancy, the system can compare each event's Count with its Quantity. This gives the percentage of tickets sold.",
            "To understand user participation, the system can group tickets by user ID. This would show which users are active participants and how many events they attend.",
            "To analyze monthly activity, the system can group events by eventDate month or tickets by createdAt month. This would show event and booking trends across time."
        ]),
        spec("3.46 Possible Admin Dashboard Components", [
            "A future admin dashboard can be divided into small reusable components. Each component should answer one clear question instead of showing too much information at once.",
            "A summary card section can show total users, total events, total tickets, total revenue, and active events. These cards would give a quick overview of the platform.",
            "A chart section can show monthly bookings, revenue by month, and event occupancy. These charts would help identify trends.",
            "A table section can show event-wise performance with columns such as event title, organizer, date, ticket price, tickets sold, capacity, revenue, and status.",
            "A verification section can show recent ticket scans, valid entries, invalid attempts, and used tickets. This would be useful on event day.",
            "A filter section can allow admins to filter by event date, organizer, price type, location, and event status. This would make the dashboard more useful when many events are stored."
        ]),
        spec("3.47 API Security Improvements", [
            "The current project includes basic authentication, but a production version would need stronger API security. Security becomes more important when payments, user data, and event records are involved.",
            "Input validation can be improved using validation libraries. This would ensure that email fields, dates, prices, and quantities follow expected formats before reaching the database.",
            "Rate limiting can be added to protect login and registration routes from repeated requests. This helps reduce brute-force attempts.",
            "Role-based middleware can be added so that only event owners can edit or delete their events and only authorized organizers can verify tickets.",
            "File upload security can be improved by checking file type, file size, and image content. Uploaded files should be stored in a secure external service in production.",
            "Audit logs can be added to record sensitive actions such as deleting an event, deleting a ticket, verifying a ticket, or changing event capacity."
        ]),
        spec("3.48 User Interface Improvements", [
            "The current interface is functional, but there are several possible improvements that would make it more polished and easier to use.",
            "The create event form can be divided into sections such as basic details, schedule, venue, ticketing, and image upload. This would make the form less overwhelming.",
            "The home page can include filters for date, free or paid events, location, and organizer. This would help users find relevant events faster.",
            "The wallet page can include sorting and filtering options such as upcoming tickets, past tickets, active tickets, cancelled tickets, and used tickets.",
            "The calendar view can show color-coded events based on category or ticket availability. This would make the schedule easier to scan.",
            "The payment and confirmation screens can be improved with clearer booking summaries, event policy reminders, and download ticket options."
        ]),
        spec("3.49 Real-World Use Case Scenario", [
            "Consider a university technical society organizing a coding workshop. The organizer logs into EventfulEMS and creates an event with the workshop title, description, date, time, venue, ticket price, and available seats.",
            "Students visiting the platform see the event card on the home page. They can open the detail page, read the description, check the time and venue, and share the link with classmates.",
            "A student books two tickets after accepting the terms. The system checks that seats are available, creates the ticket, generates the QR code, and updates the sold count.",
            "The student later opens the wallet page and sees the ticket with QR code and ticket ID. On event day, the ticket can be verified manually using the ticket code.",
            "After the event, the organizer can analyze how many tickets were sold, how many seats remained, how much revenue was generated, and whether the event received enough engagement.",
            "This scenario shows how the application supports the complete event lifecycle from creation to reporting."
        ]),
        spec("3.50 Summary of Project Implementation", [
            "EventfulEMS was implemented as a complete full stack project with frontend, backend, database, authentication, image upload, and ticket generation.",
            "The frontend provides the user-facing pages and handles navigation, forms, user state, event display, booking screens, wallet display, calendar view, and verification interface.",
            "The backend provides API routes for authentication, event management, ticket management, and health checking. It also handles image storage, token signing, password hashing, and database connection.",
            "The database stores users, events, and tickets in structured collections. These collections are connected through user ID and event ID fields.",
            "The project is analytics-ready because it stores capacity, count, price, total price, likes, ticket quantity, ticket status, and timestamps.",
            "The implementation gave me practical experience in building a system that is useful for users and meaningful from a data perspective."
        ]),
        spec("CHAPTER 4 - DESIGN, DIAGRAMS, AND SCREENSHOTS", [
            "This chapter is reserved for diagrams and screenshots. The report intentionally leaves clean blank spaces so that final application screenshots can be added after opening the project in the browser.",
            "Recommended screenshots include home page, login page, create event page, event detail page, order summary, payment confirmation, ticket wallet, my events page, calendar view, and verification center.",
            "Each image should be inserted inside the placeholder area and resized to fit neatly within the page margins."
        ]),
        spec("4.1 System Architecture Diagram", [
            "The system architecture diagram should show how the React frontend communicates with the Express API, how the API connects to MongoDB, and how uploaded images and ticket data are handled."
        ], image=("IMAGE PLACEHOLDER - SYSTEM ARCHITECTURE DIAGRAM", "Insert a clean architecture diagram showing React, Express API, MongoDB, uploads, and deployment flow.", 12)),
        spec("4.2 Data Flow Diagram", [
            "The data flow diagram should explain the movement of data from user actions to backend routes and database collections."
        ], image=("IMAGE PLACEHOLDER - DATA FLOW DIAGRAM", "Insert DFD showing user registration, event creation, ticket booking, wallet fetching, and verification.", 12)),
        spec("4.3 Database Schema Diagram", [
            "The database schema diagram should show User, Event, and Ticket entities and their relationship through user ID and event ID."
        ], image=("IMAGE PLACEHOLDER - DATABASE SCHEMA", "Insert ER-style diagram for User, Event, and Ticket collections.", 12)),
        spec("4.4 Home Page Screenshot", [
            "The home page screenshot should show event cards with images, event details, likes, price, organizer, and booking button."
        ], image=("SCREENSHOT PLACEHOLDER - HOME PAGE", "Insert screenshot of the EventfulEMS home page with multiple events visible.", 12)),
        spec("4.5 Create Event Screenshot", [
            "The create event screenshot should show the event form fields and upload option."
        ], image=("SCREENSHOT PLACEHOLDER - CREATE EVENT PAGE", "Insert screenshot of the Post an Event form.", 12)),
        spec("4.6 Event Detail Screenshot", [
            "The event detail screenshot should show the event banner, title, price, description, date, location, and sharing options."
        ], image=("SCREENSHOT PLACEHOLDER - EVENT DETAIL PAGE", "Insert screenshot of an event detail page.", 12)),
        spec("4.7 Booking Workflow Screenshot", [
            "This screenshot should show the order summary or payment summary page where ticket quantity and total price are visible."
        ], image=("SCREENSHOT PLACEHOLDER - BOOKING WORKFLOW", "Insert screenshot of order summary/payment summary page.", 12)),
        spec("4.8 Ticket Wallet Screenshot", [
            "The ticket wallet screenshot should show QR code, event name, quantity, price, email, and ticket ID."
        ], image=("SCREENSHOT PLACEHOLDER - TICKET WALLET", "Insert screenshot of booked ticket cards in the wallet.", 12)),
        spec("4.9 Calendar and Verification Screenshots", [
            "This page should include either two smaller screenshots or one combined screenshot showing the calendar view and verification center."
        ], image=("SCREENSHOT PLACEHOLDER - CALENDAR / VERIFICATION", "Insert calendar view and verification center screenshots here.", 12)),
        spec("CHAPTER 5 - TESTING AND RESULTS", [
            "Testing was done to ensure that the main workflows work correctly. Since the project is a web application, manual functional testing was performed through the browser and API behavior was checked through frontend actions.",
            "The main goal of testing was to verify that users can register, log in, create events, browse events, book tickets, view tickets, and delete records where allowed.",
            "The tests also checked that ticket quantity limits and capacity checks work correctly."
        ]),
        spec("5.1 Functional Test Cases", [
            "The following table summarizes important functional test cases used for the project."
        ], table=(["Test Case", "Expected Result", "Status"], [
            ("Register new user", "User account is created and session starts.", "Pass"),
            ("Login with valid details", "User profile loads and protected pages open.", "Pass"),
            ("Create event with image", "Event is saved and visible on listing page.", "Pass"),
            ("Book ticket", "Ticket is created, total is calculated, and count updates.", "Pass"),
            ("View wallet", "User sees booked tickets with QR code and ticket ID.", "Pass"),
            ("Delete event/ticket", "Confirmation appears and selected record is removed.", "Pass"),
        ], [Cm(5), Cm(8), Cm(3)])),
        spec("5.2 Validation and Error Handling", [
            "Validation was added at important points in the backend. Ticket creation checks whether user ID and event ID are present, whether the event exists, whether the user has exceeded the ticket limit, and whether enough tickets are available.",
            "The profile route clears invalid tokens and returns a proper error instead of letting the application fail silently.",
            "Frontend messages were added in places such as event creation, ticket deletion, and ticket booking so users can understand what happened."
        ]),
        spec("5.3 Results Achieved", [
            "The final project successfully implements the core workflow of an event management and analytics system. Events can be created, displayed, opened, shared, booked, and tracked.",
            "Ticket booking generates structured ticket records with QR code data, ticket code, event details, user details, quantity, and total price.",
            "The project also prepares a base for future analytics dashboards because it stores event capacity, sold count, price, likes, and ticket information."
        ]),
        spec("5.4 Output Screenshots", [
            "Add final tested output screenshots on this page after running the application. Prefer clear browser screenshots with the address bar hidden and the page zoom set to 100%."
        ], image=("SCREENSHOT PLACEHOLDER - TESTED OUTPUTS", "Insert final output screenshots or a collage of successful test screens here.", 14)),
        spec("CHAPTER 6 - FUTURE SCOPE AND CONCLUSION", [
            "EventfulEMS can be extended into a more advanced event analytics platform. The current implementation covers the most important workflows, but several production-level features can be added later.",
            "The most useful future addition would be a dashboard for organizers showing ticket sales, remaining capacity, revenue, likes, bookings by date, and event performance comparison.",
            "Role-based access can also be added so that admins, organizers, and attendees have separate permissions."
        ]),
        spec("6.1 Future Scope", [
            "Future enhancements can make EventfulEMS more useful for real institutions and event teams."
        ], bullets=[
            "Add a real payment gateway such as Razorpay or Stripe.",
            "Build analytics dashboards for revenue, bookings, capacity, and engagement.",
            "Add QR scanning for gate entry and mark tickets as used.",
            "Add email ticket delivery and booking reminders.",
            "Move uploaded images to Cloudinary, S3, or another persistent storage service.",
            "Add organizer approval flow and admin panel.",
            "Add filters by event category, location, date range, and price.",
            "Export event reports as CSV or PDF for organizers."
        ]),
        spec("6.2 Conclusion", [
            "The industrial training at Sabudh Foundation helped me improve my technical understanding and apply my skills to a practical project. The training exposed me to data science thinking, project-based learning, and full stack implementation.",
            "EventfulEMS became the main outcome of this training. It is a working event management and analytics-oriented system built using React, Node.js, Express, MongoDB, JWT, Multer, and QR code generation.",
            "The project helped me understand how frontend design, backend APIs, database models, authentication, and deployment planning come together in a real application. It also improved my confidence in building systems that are useful, understandable, and extendable."
        ]),
        spec("REFERENCES", [
            "1. Sabudh Foundation official website: https://sabudh.org/",
            "2. Sabudh Foundation About Us page: https://sabudh.org/about-us/",
            "3. Sabudh Foundation Apply Now page: https://sabudh.org/apply-now/",
            "4. React documentation: https://react.dev/",
            "5. Vite documentation: https://vite.dev/",
            "6. Express documentation: https://expressjs.com/",
            "7. MongoDB documentation: https://www.mongodb.com/docs/",
            "8. Mongoose documentation: https://mongoosejs.com/",
            "9. Node.js documentation: https://nodejs.org/",
            "10. Render deployment documentation: https://render.com/docs/"
        ]),
        spec("APPENDIX", [
            "Appendix A: Suggested screenshots to add before final submission.",
            "1. Home page with event cards.",
            "2. Login and registration pages.",
            "3. Create event form.",
            "4. Event details page.",
            "5. Order summary and payment summary.",
            "6. Ticket confirmation screen.",
            "7. Ticket wallet with QR code.",
            "8. My Events page.",
            "9. Calendar view.",
            "10. Verification center.",
            "Appendix B: Suggested diagrams to add.",
            "1. System architecture diagram.",
            "2. Data flow diagram.",
            "3. Database schema diagram.",
            "4. Ticket booking workflow diagram."
        ]),
    ]

    enrich = {
        "CHAPTER 1 - INTRODUCTION": [
            "The purpose of this report is to present the complete work carried out during the training period in a formal academic format. It documents the training organization, learning process, technologies explored, project design, development stages, testing methods, and outcomes.",
            "The project was selected because event management is a common real-world problem in colleges, clubs, societies, and small organizations. Most event teams need a system that can publish event information, collect bookings, manage attendance data, and make reporting easier.",
            "EventfulEMS is therefore not only a web application but also a data-driven system. Every event, ticket, price, booking count, user identity, date, and capacity value becomes a structured data point that can later support analytics and decision-making."
        ],
        "1.1 About Sabudh Foundation": [
            "Sabudh Foundation focuses on building practical capability in areas such as data science, analytics, artificial intelligence, and connected technology. The training style is not limited to theoretical lectures; it encourages students to learn by building and reviewing real project work.",
            "The foundation's learning model helped me understand the importance of approaching problems systematically. Instead of directly jumping into coding, the training encouraged identifying the problem, studying the users, choosing the right tools, and then implementing the solution in smaller parts.",
            "This approach was useful while developing EventfulEMS because the project required both software development and data thinking. I had to consider user flow, database records, validation, and the kind of analytics that could be generated from the stored data later."
        ],
        "1.2 Introduction to Internship": [
            "The internship work involved regular technical learning and project development. Since the selected major project was EventfulEMS, I focused on understanding how a full stack system can be designed in a way that also supports data analysis.",
            "The training began with revision of important concepts such as APIs, frontend routing, databases, authentication, and basic deployment. After this, I worked on converting the project idea into modules and gradually implementing them.",
            "The internship also improved my ability to document technical work. I learned that project reporting is not only about listing technologies, but also about explaining why the system was needed, how it was built, how it was tested, and how it can be improved."
        ],
        "1.3 Internship Learning Experience": [
            "One of the most important parts of the internship was learning how to debug issues independently. While working on frontend and backend integration, small configuration mistakes in API paths, cookies, environment variables, or database fields could break a complete workflow.",
            "I also learned the importance of writing features in a way that matches real user behaviour. For example, users should not be able to book unlimited tickets, create events without logging in, or proceed to payment without confirming the event details.",
            "The project improved my understanding of how software and analytics connect. A well-designed application should collect data in a clean and consistent format so that reports and dashboards can be added later without rebuilding the database."
        ],
        "1.4 Vision Language Navigation Work": [
            "This work was mainly exploratory and research-oriented. I studied how a VLN agent uses both visual input and natural language instructions to decide movement. The system must convert a human instruction into smaller navigation decisions.",
            "The important learning from VLN was that modern AI systems often combine multiple modules rather than relying on only one model. Vision models help recognize the scene, language models help understand instructions, and planning algorithms help choose a route.",
            "Although VLN is not the main project of this report, it helped me understand multimodal AI and robotics at a conceptual level. This improved my general technical awareness during the internship."
        ],
        "CHAPTER 2 - TECHNOLOGY USED": [
            "The technology stack was chosen according to the needs of the project. The system required an interactive user interface, a backend server, secure user sessions, file uploads, QR generation, and a database for storing users, events, and tickets.",
            "A MERN-based stack was suitable because it uses JavaScript across most of the application. This reduced context switching and made it easier to connect frontend actions with backend APIs.",
            "The tools used in the project were not treated as isolated technologies. Each tool had a specific role in the final workflow, from rendering event cards to validating ticket purchases and storing booking records."
        ],
        "2.1 React and Vite": [
            "React's component-based structure was useful because EventfulEMS contains many repeated interface patterns such as event cards, buttons, forms, protected screens, and confirmation sections.",
            "Vite helped during development by quickly refreshing the browser when code changed. This was useful while adjusting page layouts, checking form behaviour, and testing route navigation.",
            "The route structure also made the project easier to understand. Each page in the application represents a user task, such as browsing events, creating an event, booking tickets, checking the wallet, or verifying a ticket."
        ],
        "2.2 Tailwind CSS and Frontend Styling": [
            "Tailwind CSS made it possible to design the interface directly inside the component files. This helped in rapid development because layout, spacing, color, shadow, and responsive behaviour could be adjusted without constantly switching between files.",
            "The styling was kept simple and functional. Since this is an event management system, the interface needed to be readable, responsive, and action-oriented rather than overly decorative.",
            "Responsive classes were used so the application can work on different screen sizes. Event cards, forms, wallet tickets, and calendar cells were designed to remain usable on laptop and smaller screens."
        ],
        "2.3 Node.js and Express": [
            "Node.js allowed the backend to be written in JavaScript, which matched the frontend language. This made it easier to understand request and response handling across the complete stack.",
            "Express was used to create REST API endpoints. These endpoints form the communication layer between the frontend and the database. Whenever a user registers, logs in, creates an event, books a ticket, or deletes a record, an Express route handles the request.",
            "The backend also performs validation and data preparation. For example, ticket counts are converted into numbers, ticket limits are checked, and ticket codes are generated before saving the booking."
        ],
        "2.4 MongoDB and Mongoose": [
            "MongoDB was useful because the project data is naturally document-oriented. An event record can contain title, date, time, location, image path, ticket price, organizer details, likes, and capacity fields in a single document.",
            "Mongoose added structure to the database layer. It allowed the project to define clear schemas for users and tickets, which made records more predictable and reduced accidental inconsistent storage.",
            "The data model was designed with future analytics in mind. Event quantity, sold count, price, ticket count, ticket status, and timestamps are all fields that can be used later for reports."
        ],
        "2.5 Authentication and Security Tools": [
            "Authentication is important because the application has user-specific actions. Only logged-in users should be able to create events, view personal tickets, see their own created events, or access verification features.",
            "JWT tokens make the session lightweight. The server signs a token containing the user email and ID, and the frontend can then request profile data to know which user is active.",
            "Password hashing with bcryptjs protects user credentials. Even in a student project, this is an important practice because it follows real development standards."
        ],
        "2.6 API, Upload, and QR Tools": [
            "Axios simplified API calls from the frontend. It was used consistently for GET, POST, and DELETE operations, which made communication with the backend cleaner.",
            "Multer was necessary because event creation includes an image upload. Without upload handling, events would be limited to text-only information, which is less useful for promotion.",
            "QR generation made the ticket module more realistic. A booked ticket is not only stored as text, but also displayed visually as a scannable code, which can later support entry verification."
        ],
        "2.7 Deployment Tools": [
            "Deployment preparation was an important part of the project because a full stack project should not remain limited to local execution. The Render configuration and build scripts make the application easier to publish.",
            "The project separates environment configuration from code. This is important because database URLs and secrets should not be hardcoded into source files.",
            "The deployment approach also handles the frontend production build through the backend server. This makes hosting simpler for a student project because the API and client can run together."
        ],
        "CHAPTER 3 - INTRODUCTION TO PROJECT": [
            "The project can be understood as a bridge between event operations and event analytics. The user-facing side helps people create and book events, while the data side keeps track of information that can later be summarized and analyzed.",
            "The application does not require users to understand the backend or database. It hides the technical complexity behind a simple interface where users can browse events, click buttons, and complete tasks.",
            "From a development point of view, the project gave me experience in handling state, routes, forms, file uploads, API calls, database operations, authentication, validation, and deployment readiness."
        ],
        "3.1 Project Overview": [
            "The complete workflow starts when an organizer logs in and creates an event. The event details are stored in the database and then shown on the home page, calendar page, and my events page.",
            "An attendee can open an event, read the details, choose ticket quantity, accept terms, and confirm the booking. The booking creates a ticket record that stores the user's details and event details together.",
            "The wallet module then displays the user's tickets. This creates a digital record of participation and makes the event experience easier for both attendees and organizers."
        ],
        "3.2 Problem Statement": [
            "Manual event management often causes duplicate records, missed updates, unclear ticket availability, and difficulty in checking who has booked tickets. These problems become more visible when multiple events are running at the same time.",
            "For students and college event teams, using different tools for announcements, registrations, payment notes, and attendance lists can create confusion. A centralized system reduces this friction.",
            "The project addresses this by creating a single web application where events, users, tickets, and booking data stay connected."
        ],
        "3.3 Objectives of the Project": [
            "Another objective was to understand how a data analytics project starts from data collection. Before advanced dashboards can exist, the application must first store clean and useful data.",
            "The project also aimed to improve my practical skills in integrating multiple libraries and frameworks. It required connecting React, Express, MongoDB, JWT, Multer, QR code generation, and deployment configuration.",
            "A final objective was to create a report-ready project with clear modules, test cases, screenshots, diagrams, and future scope so that it can be submitted as industrial training work."
        ],
        "3.4 Proposed System": [
            "The proposed system follows a client-server architecture. The React client handles user interaction, while the Express server handles business logic and database communication.",
            "MongoDB acts as the storage layer for persistent data. Every important user action, such as event creation or ticket booking, results in a database update.",
            "The proposed system can later be extended into a role-based platform with separate attendee, organizer, and administrator dashboards."
        ],
        "3.5 Scope of the Project": [
            "The scope also includes basic event analytics because the project already tracks sold count, ticket quantity, likes, price, and ticket totals. These values can be visualized later using charts.",
            "The project is not limited to college events. With small changes, it can support workshops, seminars, competitions, cultural events, technical fests, meetups, or small paid events.",
            "The current scope intentionally keeps advanced features optional so that the main workflow remains stable and understandable."
        ],
        "3.6 Project Modules": [
            "Each module was designed around a specific user need. This helped keep the project organized and made it easier to test one workflow at a time.",
            "The frontend modules are connected through React Router, while the backend modules are connected through Express routes and MongoDB collections.",
            "The modular structure also makes future improvements easier because new dashboards, filters, or reports can be added without disturbing the complete application."
        ],
        "3.7 Authentication Module": [
            "The register route accepts name, email, and password. It hashes the password, creates the user document, signs a token, and sends the user information back to the frontend.",
            "The login route checks whether the user exists and whether the password matches the stored hash. If the login is valid, the system creates a token and starts a session.",
            "The profile route is used by the frontend to load the current user. This is important because several pages need to know the user ID before they can fetch personal events or tickets."
        ],
        "3.8 Event Creation Module": [
            "The event creation form includes fields that are useful for both display and analytics. Date, time, location, ticket price, and total tickets are not only shown to users but can also be used in reports.",
            "The backend converts numeric fields such as quantity and ticket price into numbers. This prevents calculation errors later when ticket count or total price is used.",
            "After successful event creation, the frontend shows a confirmation screen with options to view the event or open the user's event list."
        ],
        "3.9 Event Listing and Discovery": [
            "The listing page improves event discovery by showing only active or upcoming events. This keeps the home page relevant for users looking for current opportunities.",
            "The like feature acts as a lightweight engagement metric. While it is simple, it can still show which events are attracting attention.",
            "The capacity progress bar helps users and organizers visually understand ticket availability without reading numbers carefully."
        ],
        "3.10 Event Detail and Sharing": [
            "The event detail page is designed to give enough information before booking. A user should know the date, time, venue, organizer, description, and price before proceeding.",
            "The share features support organic promotion. Students often share events through WhatsApp or social media, so the system includes quick options for copying and sharing the event link.",
            "This module also acts as the connection point between browsing and booking because the Book Ticket button starts the ticket purchase workflow."
        ],
        "3.11 Order Summary Module": [
            "The order summary page protects the user from accidental bookings. Before moving ahead, the user can review the ticket quantity, event title, price, available seats, and terms.",
            "The quantity input is restricted so users cannot select less than one ticket or more than the allowed limit. If event capacity is low, the maximum selectable quantity is reduced automatically.",
            "This step makes the booking process more reliable and professional."
        ],
        "3.12 Payment Summary and Ticket Generation": [
            "The payment page currently simulates a card payment. This was done so the complete flow could be demonstrated without requiring a real payment gateway account.",
            "Ticket generation is the most important backend workflow. It checks event availability, prevents users from crossing the ticket limit, calculates total price, creates a unique ticket code, and saves the final ticket.",
            "After ticket creation, the event's sold count is increased. This update is important because it keeps capacity information correct for future bookings."
        ],
        "3.13 Ticket Wallet Module": [
            "The wallet is user-specific. It fetches tickets using the logged-in user's ID so that each user sees only their own bookings.",
            "The QR code and ticket code make the ticket visually and functionally identifiable. This can later be used for scanning and attendance verification.",
            "The delete option was added with a confirmation dialog because ticket deletion is a sensitive action and should not happen accidentally."
        ],
        "3.14 My Events Module": [
            "This module is useful for organizers because they need a quick view of the events they have created. It shows event title, date, time, location, price, ticket count, and capacity.",
            "The delete function allows organizers to remove events that are cancelled or no longer needed. The confirmation dialog explains that tickets already sold may still exist.",
            "In future versions, this module can become a complete organizer dashboard with edit options, attendee lists, charts, and downloadable reports."
        ],
        "3.15 Calendar View Module": [
            "The calendar module uses date logic to generate all days of the current month and place events on matching dates. This gives a planner-style view of the event schedule.",
            "Users can move between months using previous and next buttons. This is useful for checking upcoming events in the near future.",
            "For institutions, a calendar view is especially helpful because events often need to be planned around exams, holidays, workshops, and other academic activities."
        ],
        "3.16 Verification Center": [
            "The verification center currently checks a ticket ID entered manually. If the entered value matches one of the user's ticket codes, the system displays a verified result.",
            "This module proves the basic logic required for ticket validation. In future, the same idea can be extended into QR scanning where an organizer scans a ticket and the backend updates the ticket status.",
            "A production version would need stronger access control so that only authorized organizers or administrators can verify tickets for a particular event."
        ],
        "3.17 Data Analytics Perspective": [
            "The analytics value of EventfulEMS comes from the relationship between event records and ticket records. Each booking links a user to an event with date, quantity, price, and ticket status.",
            "This data can be used to build descriptive analytics such as total bookings, revenue, occupancy percentage, popular events, and monthly event activity.",
            "It can also support decision-making. For example, if paid events have low bookings but free events have high engagement, organizers can adjust pricing or promotion strategy."
        ],
        "3.18 Important Data Entities": [
            "The User entity is simple but important because it identifies who is creating events and who is booking tickets. User identity is the base for personalization.",
            "The Event entity is the core of the system. It stores all public-facing event information as well as analytics fields such as likes, quantity, and sold count.",
            "The Ticket entity works like a transaction record. It stores the actual booking and therefore becomes the most useful collection for revenue and attendance analysis."
        ],
        "3.19 Capacity and Booking Analytics": [
            "Capacity percentage can be calculated by dividing sold tickets by total ticket quantity. This gives a quick measure of how full an event is.",
            "Booking analytics can also help detect demand. Events that sell out quickly may need a larger venue or additional sessions, while low-booked events may need better promotion.",
            "This kind of analytics is valuable because it turns raw booking records into planning information."
        ],
        "3.20 Revenue and Price Analytics": [
            "Revenue analytics can be calculated at event level, monthly level, or organizer level. The ticket records already store total price, so the system has the basic data needed for this.",
            "Different event types can be compared by price and booking count. This can help organizers decide whether an event should be free, low-cost, or paid.",
            "For future development, a dashboard can display total revenue, average booking value, top earning events, and refund/cancellation trends."
        ],
        "3.21 User Engagement Analytics": [
            "Engagement analytics can combine likes, bookings, and sharing behaviour. Even if sharing is not stored in the current version, the application design already includes share actions that can later be tracked.",
            "An event with many likes but fewer bookings may indicate interest without conversion. This may happen due to high price, inconvenient timing, or unclear details.",
            "Such insights are useful because they help improve event planning beyond simple attendance counting."
        ],
        "3.22 Backend API Design": [
            "The API layer is the backbone of the project. It ensures that frontend actions do not directly access the database and instead go through controlled backend routes.",
            "Using API routes also makes validation easier. The backend can reject invalid requests before they affect the database.",
            "The current API design is simple enough for the project but can be expanded into separate route files, controller files, middleware, and services in a larger version."
        ],
        "3.23 Database Design": [
            "The database design focuses on keeping related values together. For example, ticketDetails inside a ticket record stores name, email, event name, event date, event time, price, total price, and QR code.",
            "This makes the ticket self-contained. Even if event details change later, the ticket still stores the details that were valid when the user booked it.",
            "Timestamps in the ticket schema also support future analytics such as bookings per day, peak booking times, and booking history."
        ],
        "3.24 Frontend Architecture": [
            "The frontend uses a clear separation between page components and utility files. Pages handle visible screens, while utility files handle API and media URL logic.",
            "UserContext improves the structure because authentication state is needed in many places. Without context, each protected page would need repeated code for loading user information.",
            "The frontend architecture can be extended by adding reusable form components, dashboard components, and chart components in future versions."
        ],
        "3.25 Backend Architecture": [
            "The backend starts by checking important environment variables. If MongoDB URL or JWT secret is missing, the server stops early instead of failing later in an unclear way.",
            "The server also creates the uploads directory automatically if it does not exist. This prevents image upload errors during first-time setup.",
            "Static serving of uploads and production client build makes the backend responsible for both data and final deployment hosting."
        ],
        "3.26 Implementation Challenges": [
            "Another challenge was keeping the user experience smooth after actions. For example, after event creation the user should see a success page, and after ticket creation the user should receive a clear ticket code.",
            "Ticket validation required careful thinking because it combines multiple conditions: user must be present, event must exist, ticket count must be valid, personal limit must not be crossed, and capacity must remain available.",
            "Handling images also required consistency between database image path, backend static route, and frontend rendering URL."
        ],
        "3.27 Solutions Applied": [
            "The project uses confirmation dialogs for delete actions. This avoids accidental deletion and improves user confidence while using the application.",
            "Error and success messages were added to make the system more transparent. Users should know whether an event was created, a ticket was deleted, or a booking failed.",
            "The use of environment variables and Render build scripts made the project easier to configure for deployment."
        ],
        "3.28 Testing Approach": [
            "Testing was done from both attendee and organizer perspectives. As an attendee, I checked browsing, event detail viewing, ticket booking, wallet display, and ticket deletion.",
            "As an organizer, I checked event creation, my events listing, event deletion, and capacity changes after bookings.",
            "Testing also included checking protected pages by opening them without login to make sure the user is redirected properly."
        ],
        "3.29 Benefits of the Project": [
            "The project improves organization because event and ticket data are stored digitally rather than being scattered across messages or spreadsheets.",
            "It improves user convenience because attendees can discover events and store tickets in one place.",
            "It improves decision-making potential because the stored data can later be used for dashboards, reports, and event planning."
        ],
        "3.30 Limitations of Current Version": [
            "The current project does not yet include an admin dashboard. Admin-level controls would be needed in a real institution for approving events and monitoring platform activity.",
            "The project does not yet include event editing after creation. This would be useful when organizers need to update venue, time, description, or image.",
            "The analytics are currently discussed and partially supported by data fields, but chart-based dashboards are not yet implemented."
        ],
        "CHAPTER 4 - DESIGN, DIAGRAMS, AND SCREENSHOTS": [
            "The diagrams in this chapter should be inserted after final browser testing. A good screenshot section improves the report because it proves that the project was implemented and not only described theoretically.",
            "The screenshots should be clean and consistent. The browser should be zoomed to 100%, unnecessary tabs should be hidden if possible, and each screenshot should show the important part of the page.",
            "Diagrams should be simple and readable. They should show data movement and system structure clearly rather than using too many symbols."
        ],
        "CHAPTER 5 - TESTING AND RESULTS": [
            "Testing is important because a project may look complete but still fail during actual use. EventfulEMS includes several connected workflows, so each workflow had to be checked carefully.",
            "The main testing goal was to verify that user actions produce correct database changes and that database records display properly on the frontend.",
            "Special attention was given to ticket booking because it affects both the Ticket collection and the Event sold count."
        ],
        "5.1 Functional Test Cases": [
            "The functional test cases were selected according to the main user journeys in the system. These journeys cover authentication, event creation, event browsing, ticket booking, wallet management, and deletion.",
            "Each test case was checked manually through the application interface. The expected result was compared with the actual screen response and database-related behaviour.",
            "The table below summarizes the most important tests performed during project validation."
        ],
        "5.2 Validation and Error Handling": [
            "Validation is necessary because users may submit incomplete forms, select invalid ticket quantities, or try to access pages without logging in.",
            "The backend performs important validation during ticket creation because this is where data consistency matters most. If invalid data entered the ticket collection, analytics and wallet display would become unreliable.",
            "Frontend messages also improve usability because users get immediate feedback instead of wondering whether an action worked."
        ],
        "5.3 Results Achieved": [
            "The completed project demonstrates the main features required in an event management platform. It supports user accounts, event creation, event browsing, booking, ticket generation, and wallet management.",
            "The data model is strong enough to support future analytics. Event quantity, sold count, likes, ticket count, total price, and timestamps can all be used to build reports.",
            "The final result is suitable for academic submission because it combines practical implementation with clear scope for future improvement."
        ],
        "CHAPTER 6 - FUTURE SCOPE AND CONCLUSION": [
            "The future scope of the project is strong because the current foundation already stores the right kind of data. The next step is to convert this stored data into visual reports and dashboards.",
            "Another major future improvement is role-based access. Separating attendees, organizers, and administrators would make the application more suitable for real deployment.",
            "The project can also be improved by adding notifications, real payment handling, better verification, and persistent image storage."
        ],
        "6.1 Future Scope": [
            "A future analytics dashboard can show cards and charts for total events, total bookings, total revenue, available seats, top events, and monthly trends.",
            "A QR scanner can be added using the device camera. When a ticket is scanned, the backend can verify the ticket code and mark the ticket as used.",
            "The project can also include email automation so users receive ticket confirmation messages and event reminders."
        ],
        "6.2 Conclusion": [
            "The project strengthened my confidence in building full stack applications. I learned how small modules combine into a complete system and how backend validation protects the quality of stored data.",
            "The training also helped me understand the importance of documentation. A good project report should explain the problem, solution, technology, implementation, testing, and future scope in a clear manner.",
            "Overall, the internship and project work improved my technical knowledge, problem-solving ability, and readiness to work on practical software and analytics projects."
        ],
        "REFERENCES": [
            "Additional references used for technical understanding include official documentation for the libraries and frameworks used in the application.",
            "The EventfulEMS source code present in the project folder was used as the primary reference for implementation details.",
            "The Vision Language Navigation content was summarized from the provided VLN presentation."
        ],
        "APPENDIX": [
            "Appendix C: Suggested analytics to implement later.",
            "1. Event occupancy percentage.",
            "2. Revenue per event.",
            "3. Bookings by month.",
            "4. Most liked events.",
            "5. Paid versus free event comparison.",
            "6. User booking history.",
            "7. Ticket status report.",
            "Appendix D: Suggested database improvements.",
            "1. Add event category.",
            "2. Add organizer role.",
            "3. Add ticket used timestamp.",
            "4. Add payment transaction ID.",
            "5. Add cancellation reason."
        ],
    }

    for idx, item in enumerate(pages):
        title, paragraphs, kwargs = item
        if title in enrich:
            pages[idx] = (title, paragraphs + enrich[title], kwargs)

    for title, paragraphs, *rest in pages:
        kwargs = {}
        if rest:
            maybe = rest[0]
            if isinstance(maybe, dict):
                kwargs = maybe
        if kwargs.get("image"):
            kwargs["break_after"] = True
        add_page(doc, title, paragraphs, **kwargs)

    doc.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT.resolve())

